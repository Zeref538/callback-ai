"""Pull plain text out of an uploaded resume / job post file.

Handles the formats people actually upload -- PDF, Word, plain text/markdown --
so the candidate can drop a file instead of copy-pasting. Digital documents
only: a scanned/photographed PDF has no text layer to extract, and true OCR
needs a system Tesseract install, so those are detected and reported rather
than silently returning nothing.
"""
import io

from pypdf import PdfReader
import docx


class UnsupportedDocument(ValueError):
    """The file type isn't one we can read, or a PDF was scanned (no text layer)."""


def _from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if len(text) < 20:
        # A resume/job post with almost no extractable text is a scanned image.
        raise UnsupportedDocument(
            "This PDF looks scanned (no selectable text). Paste the text instead, "
            "or export a text-based PDF."
        )
    return text


def _from_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" \t".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p.strip()).strip()


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on extension. filename is only used for its suffix."""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith((".txt", ".md", ".text")):
        return data.decode("utf-8", errors="replace").strip()
    if name.endswith(".doc"):
        raise UnsupportedDocument("Legacy .doc isn't supported — save as .docx or paste the text.")
    raise UnsupportedDocument(f"Can't read {filename!r}. Upload a PDF, DOCX, TXT, or MD file.")
