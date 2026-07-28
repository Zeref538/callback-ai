"""Unit tests for document text extraction (resume/job-post uploads)."""
import io

import docx
import pytest

from callback_ai.ingest.document import UnsupportedDocument, extract_text


def _make_docx(paragraphs, table_rows=None) -> bytes:
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        t = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def test_txt_and_md_and_text_extensions():
    assert extract_text("resume.txt", b"Hello world") == "Hello world"
    assert extract_text("notes.md", b"# Title\n\nBody") == "# Title\n\nBody"
    assert extract_text("a.text", b"  trimmed  ") == "trimmed"


def test_txt_bad_bytes_dont_crash():
    # invalid utf-8 is replaced, not raised
    out = extract_text("resume.txt", b"caf\xe9 role")
    assert "role" in out


def test_docx_reads_paragraphs_and_tables():
    data = _make_docx(
        ["Jane Doe — Backend Engineer", "Built a Redis retry layer."],
        table_rows=[["Skill", "Level"], ["Python", "Expert"]],
    )
    text = extract_text("resume.docx", data)
    assert "Redis retry layer" in text
    assert "Python" in text and "Expert" in text


def test_legacy_doc_rejected():
    with pytest.raises(UnsupportedDocument):
        extract_text("old.doc", b"\xd0\xcf\x11\xe0stuff")


def test_unknown_extension_rejected():
    with pytest.raises(UnsupportedDocument):
        extract_text("archive.zip", b"PK\x03\x04")


def test_empty_filename_rejected():
    with pytest.raises(UnsupportedDocument):
        extract_text("", b"whatever")


def test_scanned_pdf_detected():
    # A PDF whose pages yield <20 chars of text is treated as scanned.
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)  # no text layer
    buf = io.BytesIO()
    writer.write(buf)
    with pytest.raises(UnsupportedDocument):
        extract_text("scan.pdf", buf.getvalue())
